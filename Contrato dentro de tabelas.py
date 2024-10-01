import pandas as pd
from datetime import datetime
from docx import Document

def substituir_texto_em_tabela(docx_file, referencias, dados):
    doc = Document(docx_file)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in referencias.items():
                        if key in paragraph.text:
                            if value in ['Inicio', 'Fim']:
                                # Certifique-se de que dados[value] é um objeto datetime
                                if isinstance(dados[value], str):
                                    # Converta a string para datetime
                                    dados[value] = datetime.strptime(dados[value].strip(),
                                                                     '%d/%m/%Y')  # ajuste o formato da data conforme necessário

                                # Usando strftime
                                data_formatada = dados[value].strftime('%b %d')
                                paragraph.text = paragraph.text.replace(key, data_formatada)
                            else:
                                paragraph.text = paragraph.text.replace(key, str(dados[value]))

    for paragraph in doc.paragraphs:
        for key, value in referencias.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(dados[value]))
    doc.save(fr"C:caminho de pasta\Contrato\Contratos feitos\ {nome_arquivo}")
    return(doc)

# Leitura dos dados do Excel
dados = pd.read_excel('DadosMSC.xlsx')

# Mapeamento das referências
referencias = {
    "VVVV": 'Vinte',
    "QQQQ": 'Quarenta',
    "OOOO": 'POL',
    "PPPP": 'POD',
    "NNNN": 'Navio',
    "YYYY": 'Viagem',
    "IIII": 'Inicio',
    "FFFF": 'Fim',
    "RRRR": 'Ano',
    "DDDD": 'Data',
    "TTTT": 'Taxa'
}

# Mapeamento dos contratos das empresas
contratos_empresas = {
    "MSC": r'C:caminho de pasta\MODELO CONTRATO MSC.docx',
}

# Processamento de todas as linhas dos dados
for index, row in dados.iterrows():
    data_atual = datetime.now()
    row['Data'] = data_atual.strftime("%d de %B de %Y")
    nome_arquivo = ("Contrato " + ' ' + row['Navio'] + ' VG. ' + str(row['Viagem']) + ' - ' + row['POL'] + '.docx')
    # Chamada da função para substituir texto em tabelas
    doc = substituir_texto_em_tabela(contratos_empresas[row['Empresa']], referencias, row)
